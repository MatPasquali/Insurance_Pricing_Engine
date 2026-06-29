"""Gera a documentacao do projeto em Word (.docx).

Fonte canonica da documentacao = este script. A cada avanco do projeto:
  1. atualize o conteudo / a constante VERSION / a lista CHANGELOG abaixo;
  2. rode  ->  python docs/gerar_documentacao.py
  3. o arquivo docs/Documentacao_Pricing_Engine.docx e regenerado.

Tambem da para editar o .docx direto no Word; nesse caso, avise para dobrarmos
as mudancas de volta neste script (a proxima regeneracao sobrescreve o arquivo).
"""
from __future__ import annotations

import pathlib

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

VERSION = "0.2.0"
LAST_UPDATE = "2026-06-13"
STATUS = "v2 CONCLUIDA (elasticidade + otimizacao de premio na base ES; v1 completa)"

OUT = pathlib.Path(__file__).resolve().parent / "Documentacao_Pricing_Engine.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)  # azul corporativo

# (versao, data, [mudancas])
CHANGELOG = [
    ("0.2.0", "2026-06-13", [
        "v2 CONCLUIDA: base ES (es_data.py/es_models.py, run_pricing_optimization.py, notebook 04).",
        "Custo GBM-Poisson (capado) + retencao logistica/GBM (AUC ~0,70 out-of-time) + otimizacao lucro x retencao.",
        "Validacao temporal (treino <=2016, teste 2017). Caveat de endogeneidade: otimo irrestrito (+71%) e red flag; restrito +-15% (+28%) e o crivel.",
    ]),
    ("0.1.7", "2026-06-13", [
        "Demo Streamlit (streamlit_app.py): perfil -> premio puro (GLM e GBM) + explicacao SHAP da previsao.",
        "v1 CONCLUIDA: dados -> modelo -> explicacao -> diagnostico -> demo interativa.",
    ]),
    ("0.1.6", "2026-06-12", [
        "Notebook 00_overview.ipynb: indice navegavel (pitch, links, resultados em destaque, roadmap).",
    ]),
    ("0.1.5", "2026-06-12", [
        "Notebook 03_diagnostics.ipynb (executado): relatividades do GLM, observado vs previsto por variavel, lift chart e Gini/Lorenz (GLM vs GBM).",
        "Validacao temporal registrada como meta do v2 (freMTPL2 nao tem dimensao temporal).",
    ]),
    ("0.1.4", "2026-06-12", [
        "Notebook 02_modeling.ipynb (executado): narrativa GLM freq/sev -> premio puro -> GLM vs GBM + SHAP.",
        "Narrativa do v1 completa (falta apenas a demo Streamlit para fechar a fase).",
    ]),
    ("0.1.3", "2026-06-12", [
        "ETL (run_etl.py): snapshots processados em parquet em data/processed (camada silver/gold).",
        "Notebook 01_eda.ipynb: EDA narrativa (frequencia rara, severidade cauda pesada, drivers BonusMalus/idade).",
        "Separacao explicita: engine (src/ + run_*) vs narrativa (notebooks/).",
    ]),
    ("0.1.2", "2026-06-09", [
        "GBM challenger (HistGradientBoosting, loss Poisson): deviance +7,97% vs ingenuo, batendo o GLM em +5,15%.",
        "Camada SHAP (TreeExplainer) explica o GBM; coeficientes do GLM e SHAP convergem nos mesmos drivers.",
        "Conclusao: ganho de acuracia confiavel, nao espurio (BonusMalus, VehAge, DrivAge). Figura SHAP embutida.",
    ]),
    ("0.1.1", "2026-06-08", [
        "Script download_datasets.py: materializa o freMTPL2 em data/raw (CSV) e",
        "baixa as bases ES (v2) e BR (v4) do Kaggle quando ha credencial da API.",
        "Documentado onde ficam as bases: cache OpenML (~/scikit_learn_data) vs data/raw.",
    ]),
    ("0.1.0", "2026-06-08", [
        "Decisao de estrategia: flagship modular (nao trilha larga).",
        "Scaffold + git; camada de dados (freMTPL2 via OpenML, sem auth).",
        "Frequencia (Poisson GLM): deviance -2,97% vs media ingenua em 678k apolices.",
        "Severidade (Gamma GLM) + premio puro (freq x sev e Tweedie direto).",
        "Achado: freq x severidade descalibra a carteira em +33%; Tweedie direto calibra a 1,04.",
        "Documentacao Word gerada por script versionado.",
    ]),
]


# --------------------------------------------------------------------------- #
# Helpers de formatacao
# --------------------------------------------------------------------------- #
def add_table(doc, headers, rows, style="Light Grid Accent 1"):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = style
    except KeyError:
        table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


# --------------------------------------------------------------------------- #
# Construcao do documento
# --------------------------------------------------------------------------- #
def build() -> None:
    doc = Document()
    doc.core_properties.title = "Insurance Pricing Engine - Documentacao"
    doc.core_properties.author = "Mateus de Pasquali da Silva"

    normal = doc.styles["Normal"].font
    normal.name = "Calibri"
    normal.size = Pt(11)

    # --- Capa ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Insurance Pricing Engine")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Motor de pricing de seguros - GLM (frequencia x severidade) + Explicabilidade (XAI)")
    rs.italic = True
    rs.font.size = Pt(13)

    doc.add_paragraph()
    meta = add_table(
        doc, ["Campo", "Valor"],
        [
            ["Autor", "Mateus de Pasquali da Silva"],
            ["Posicionamento", "Data Science aplicado a Pricing, Portfolio e Credit Risk (seguros)"],
            ["Versao do documento", VERSION],
            ["Ultima atualizacao", LAST_UPDATE],
            ["Status", STATUS],
            ["Repositorio", "insurance-pricing-engine/ (git local, sem remoto ainda)"],
        ],
    )
    doc.add_page_break()

    # --- 1. Sumario executivo ---
    doc.add_heading("1. Sumario executivo", level=1)
    add_para(doc,
        "Este projeto constroi um motor de pricing de seguros pelo metodo atuarial classico "
        "(frequencia x severidade -> premio puro), acrescido de uma camada de explicabilidade "
        "(SHAP) que compara um GLM - interpretavel por construcao - a um modelo de gradient "
        "boosting mais preciso. O objetivo e demonstrar, num contexto regulado, o dominio do "
        "trade-off precisao x interpretabilidade, cruzando a linha de pesquisa de XAI (mestrado) "
        "com a atuacao profissional em pricing.")
    add_para(doc,
        "E o projeto flagship do portfolio: um unico repositorio que serve de espinha dorsal e "
        "absorve as demais frentes (otimizacao de premio, MLOps e score territorial brasileiro) "
        "como modulos do mesmo projeto.")

    # --- 2. Objetivo e contexto ---
    doc.add_heading("2. Objetivo e contexto de portfolio", level=1)
    add_para(doc,
        "Lacunas que o projeto fecha para vagas de pricing/risco: modelagem preditiva regulada, "
        "otimizacao aplicada e deploy/MLOps. Diferencial a explorar: XAI (mestrado) x pricing "
        "(profissao). A narrativa de negocio do README e bilingue (PT/EN).")
    doc.add_heading("Problema de negocio", level=2)
    add_para(doc,
        "Uma seguradora precisa cobrar de cada apolice um premio que reflita o risco esperado - "
        "nem alto a ponto de perder o cliente, nem baixo a ponto de dar prejuizo. O custo esperado "
        "de sinistros decompoe-se em:")
    p = doc.add_paragraph()
    pr = p.add_run("Premio puro  =  Frequencia (sinistros/ano)  x  Severidade (custo medio por sinistro)")
    pr.bold = True

    # --- 3. Estrategia ---
    doc.add_heading("3. Decisao de estrategia", level=1)
    add_para(doc,
        "Decisao (2026-06-08): aprofundar um flagship modular, e nao montar uma trilha larga de "
        "varios projetos rasos. Razoes:")
    add_bullets(doc, [
        "Tempo: emprego full-time + mestrado em andamento; um repositorio concentra o esforco.",
        "Convergencia de gaps: modelagem regulada, otimizacao e MLOps cabem num unico projeto de pricing.",
        "Diferencial unico: XAI x pricing some numa trilha generica; brilha num flagship.",
        "Recrutador: clica em um link com demo viva e README de negocio, nao monta a trilha na cabeca.",
    ])
    add_para(doc, "Frase de elevador:", bold=True)
    add_para(doc,
        "\"Construi um sistema de pricing de seguros ponta a ponta: nucleo GLM regulado e "
        "explicavel, camada de otimizacao de premio por elasticidade, produtizado com "
        "monitoramento de drift.\"", italic=True)

    # --- 4. Datasets ---
    doc.add_heading("4. Dados", level=1)
    add_table(
        doc, ["Dataset", "Uso", "Acesso"],
        [
            ["freMTPL2 (French Motor TPL)", "v1 - primario",
             "OpenML fetch_openml 41214/41215 (sem auth). Espelho Kaggle: karansarpal/..."],
            ["Motor Vehicle Insurance Portfolio (ES, ~105k, 2015-2018)", "v2 - preco/retencao reais",
             "Kaggle mexwell/... = jocelyndumlao/... (espelhos); tem Premium + Lapse + renovacoes"],
            ["Brazilian Motor Insurance Market", "v4 - territorial BR",
             "Kaggle rodrigodomingos/...; provavel SUSEP/Autoseg; casa com IBGE/CNEFE/CEP"],
        ],
    )
    add_para(doc,
        "Observacao: o portfolio espanhol traz preco cobrado e churn reais, eliminando a "
        "necessidade de simular demanda no v2 (e habilitando analise de sobrevivencia de lapso).")
    doc.add_heading("Onde ficam as bases", level=2)
    add_bullets(doc, [
        "freMTPL2: baixado via OpenML e cacheado em ~/scikit_learn_data; rode "
        "'python download_datasets.py' para materializar copias em data/raw (CSV, ~35 MB).",
        "ES (v2) e BR (v4): exigem credencial da API do Kaggle (kaggle.json em ~/.kaggle/) "
        "e sao baixadas pelo mesmo script para data/raw.",
        "data/raw e data/processed sao gitignored (dado nao vai para o repositorio).",
    ])

    # --- 5. Metodologia ---
    doc.add_heading("5. Metodologia", level=1)
    add_table(
        doc, ["Componente", "Modelo", "Distribuicao", "Observacao"],
        [
            ["Frequencia", "GLM (PoissonRegressor)", "Poisson / Neg. Binomial",
             "Alvo = ClaimNb/Exposure; Exposure como peso (offset de exposicao)"],
            ["Severidade", "GLM (GammaRegressor)", "Gamma",
             "Apenas apolices com sinistro; alvo = ClaimAmount/ClaimNb; peso = ClaimNb"],
            ["Premio puro (produto)", "Frequencia x Severidade", "-",
             "Multiplica os dois GLMs"],
            ["Premio puro (direto)", "GLM (TweedieRegressor)", "Tweedie (power=1,9)",
             "Compound Poisson-Gamma: trata zeros e cauda num modelo so"],
            ["Explicabilidade", "SHAP: GLM vs GBM", "-", "v1 - concluido (GBM bate GLM em +5,15%)"],
        ],
    )

    # --- 6. Resultados ---
    doc.add_heading("6. Resultados atuais (v1)", level=1)
    doc.add_heading("6.1 Frequencia", level=2)
    add_para(doc, "GLM Poisson sobre 678.013 apolices; frequencia media de 0,1005 sinistros/ano.")
    add_table(
        doc, ["Modelo", "Mean Poisson Deviance (test)"],
        [["Media ingenua", "0,62738"], ["GLM Poisson", "0,60874  (-2,97%)"]],
    )
    doc.add_heading("6.2 Premio puro", level=2)
    add_para(doc,
        "Avaliacao em 135.603 apolices de teste (mean Tweedie deviance, p=1,9) e calibracao de "
        "carteira (total previsto / total real; sinistro real no teste = 11.868.104 EUR).")
    add_table(
        doc, ["Abordagem", "Tweedie Deviance", "Total prev/real"],
        [
            ["Media ingenua", "34,80", "1,01"],
            ["Frequencia x Severidade (Gamma)", "34,69", "1,33"],
            ["Tweedie direto", "34,46", "1,04"],
        ],
    )
    add_para(doc,
        "Insight: o modelo frequencia x severidade ranqueia o risco, mas descalibra o total da "
        "carteira em +33% (cauda pesada da severidade no freMTPL2). O Tweedie direto vence em "
        "deviance e fica quase perfeitamente calibrado (1,04). Esse e o tipo de discussao "
        "precisao x calibracao que demonstra maturidade de pricing.", bold=False)
    doc.add_heading("6.3 Explicabilidade: GLM vs GBM (XAI)", level=2)
    add_para(doc,
        "Desafiante GBM (HistGradientBoosting, loss Poisson) na mesma base de frequencia:")
    add_table(
        doc, ["Modelo", "Mean Poisson Deviance", "vs ingenuo"],
        [
            ["GLM (Poisson)", "0,60874", "+2,97%"],
            ["GBM (Poisson)", "0,57737", "+7,97%"],
        ],
    )
    add_para(doc,
        "O GBM ganha +5,15% de deviance sobre o GLM ao capturar nao-linearidades e interacoes. "
        "A camada SHAP mostra que ele se apoia nos mesmos drivers atuarialmente coerentes que o "
        "GLM - BonusMalus, idade do veiculo e idade do condutor - logo o ganho de acuracia e "
        "confiavel, nao espurio. Esse e o trade-off precisao x interpretabilidade que pricing "
        "regulado precisa pesar.")
    fig = pathlib.Path(__file__).resolve().parent.parent / "reports" / "figures" / "shap_gbm_frequency.png"
    if fig.exists():
        doc.add_picture(str(fig), width=Inches(5.5))
    add_para(doc,
        "Baselines propositalmente enxutos (sem feature engineering ainda); splines e interacoes "
        "estreitam a vantagem do GBM e reforcam a leitura via SHAP.", italic=True)

    # --- 7. Arquitetura ---
    doc.add_heading("7. Arquitetura do projeto", level=1)
    code = doc.add_paragraph()
    cr = code.add_run(
        "insurance-pricing-engine/\n"
        "  src/pricing/\n"
        "    data.py          # loaders freMTPL2 (OpenML) + frame de modelagem\n"
        "    features.py      # pre-processamento compartilhado (one-hot + scaler)\n"
        "    frequency.py     # Poisson GLM\n"
        "    severity.py      # Gamma GLM\n"
        "    pure_premium.py  # Tweedie GLM\n"
        "    gbm.py           # HistGradientBoosting (challenger)\n"
        "    etl.py           # raw -> processed (parquet)\n"
        "  run_etl.py                  # gera data/processed\n"
        "  run_frequency_baseline.py   # baseline de frequencia\n"
        "  run_pure_premium.py         # comparacao de premio puro\n"
        "  run_gbm_vs_glm.py           # GLM vs GBM + SHAP\n"
        "  download_datasets.py        # materializa as bases em data/raw\n"
        "  streamlit_app.py            # demo interativa (perfil -> premio + SHAP)\n"
        "  notebooks/00_overview.ipynb # indice navegavel\n"
        "  notebooks/01_eda.ipynb      # narrativa: EDA executada\n"
        "  notebooks/02_modeling.ipynb # narrativa: modelos + SHAP\n"
        "  notebooks/03_diagnostics.ipynb # viz: relatividades, lift, Gini\n"
        "  docs/                       # esta documentacao (script + .docx)\n"
        "  reports/figures/  tests/  data/raw + data/processed (gitignored)\n"
        "  README.md  requirements.txt"
    )
    cr.font.name = "Consolas"
    cr.font.size = Pt(9)

    # --- 8. Como rodar ---
    doc.add_heading("8. Como rodar", level=1)
    run_code = doc.add_paragraph()
    rc = run_code.add_run(
        "pip install -r requirements.txt\n"
        "python run_frequency_baseline.py\n"
        "python run_pure_premium.py"
    )
    rc.font.name = "Consolas"
    rc.font.size = Pt(10)
    add_para(doc,
        "Ambiente: Python 3.14; scikit-learn 1.8, shap, pandas, numpy, matplotlib. O dado e "
        "baixado do OpenML na primeira execucao (cache local) - nao precisa de credencial Kaggle.")

    # --- 9. Roadmap ---
    doc.add_heading("9. Roadmap (flagship modular)", level=1)
    add_table(
        doc, ["Fase", "Conteudo", "Status"],
        [
            ["v1", "Nucleo GLM + XAI: freq, sev, premio puro, GBM+SHAP, notebooks, demo Streamlit", "Concluida"],
            ["v2", "Elasticidade e otimizacao de premio (portfolio ES, dados reais de preco/lapso)", "Concluida"],
            ["v3", "MLOps: MLflow, Docker, drift PSI, CI no GitHub Actions", "Planejado"],
            ["v4", "Score territorial Brasil (IBGE/CNEFE/CEP + SUSEP/Autoseg) como feature", "Planejado"],
        ],
    )

    # --- 10. Proximos passos ---
    doc.add_heading("10. Proximos passos", level=1)
    add_bullets(doc, [
        "Feature engineering: splines em idade/densidade e interacoes (estreitar vantagem do GBM).",
        "Estender o SHAP ao modelo de severidade e ao premio puro.",
        "Demo Streamlit (perfil -> premio + explicacao) e push para o GitHub.",
        "Depois: v2 elasticidade, v3 MLOps, v4 territorial BR.",
    ])

    # --- 11. Changelog ---
    doc.add_heading("11. Historico de versoes (changelog)", level=1)
    for ver, date, changes in CHANGELOG:
        doc.add_heading(f"v{ver} - {date}", level=2)
        add_bullets(doc, changes)

    doc.save(OUT)
    print(f"Documento gerado: {OUT}  (v{VERSION})")


if __name__ == "__main__":
    build()
